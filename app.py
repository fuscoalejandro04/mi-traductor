import streamlit as st
import io
import fitz  # PyMuPDF
import re
import os
import time
import json
import hashlib
from concurrent.futures import ThreadPoolExecutor, as_completed
from deep_translator import GoogleTranslator
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_LEFT
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    SimpleDocTemplate, BaseDocTemplate, PageTemplate, Frame, Paragraph, Spacer, 
    PageBreak, NextPageTemplate
)
from reportlab.platypus.tableofcontents import TableOfContents
from xml.sax.saxutils import escape
from datetime import datetime

# Librería de Exportación: Word
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ============================================================
# CONFIGURACIÓN Y CONSTANTES
# ============================================================
MAX_CHARS_PER_FRAGMENT = 4500
FONT_PATH = os.path.join(os.path.dirname(__file__), "fonts", "DejaVuSans.ttf")
CHECKPOINT_DIR = os.path.join(os.path.dirname(__file__), ".checkpoints")

WORKERS_PARALELOS = 2
TAMANO_LOTE = 15
PAUSA_BASE = 3.0
MAX_PAUSA = 15.0

if not os.path.exists(CHECKPOINT_DIR):
    os.makedirs(CHECKPOINT_DIR)

# ============================================================
# GESTIÓN DE CHECKPOINTS ATÓMICOS EN DISCO
# ============================================================
def generar_hash_archivo(pdf_bytes):
    return hashlib.md5(pdf_bytes).hexdigest()

def cargar_checkpoint(file_hash):
    filepath = os.path.join(CHECKPOINT_DIR, f"{file_hash}.json")
    if os.path.exists(filepath):
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                return json.load(f)
        except json.JSONDecodeError:
            pass 
            
    return {
        "archivo_hash": file_hash,
        "procesados": 0,
        "resultados": {}, 
        "metricas": {
            "requests_enviados": 0,
            "reintentos": 0,
            "fallbacks": 0,
            "tiempo_acumulado": 0.0
        }
    }

def guardar_checkpoint(datos, file_hash):
    filepath = os.path.join(CHECKPOINT_DIR, f"{file_hash}.json")
    tmp_path = filepath + ".tmp"
    with open(tmp_path, 'w', encoding='utf-8') as f:
        json.dump(datos, f, ensure_ascii=False, indent=2)
    os.replace(tmp_path, filepath)

# ============================================================
# PREPROCESADOR Y FILTRO INTELIGENTE
# ============================================================
def necesita_traduccion(texto):
    if not texto or len(texto) < 2: return False
    if re.match(r'^[\d\s\W_]+$', texto): return False 
    if texto.startswith("http") and " " not in texto: return False 
    return True

# ============================================================
# MOTOR DE EXTRACCIÓN (PyMuPDF)
# ============================================================
def detectar_estructura_pymupdf(pdf_bytes):
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    bloques_estructurados = []
    tamanos_fuentes = {}
    
    for pagina in doc:
        bloques = pagina.get_text("dict")["blocks"]
        for b in bloques:
            if b["type"] == 0:
                for linea in b["lines"]:
                    for span in linea["spans"]:
                        size = round(span["size"], 1)
                        tamanos_fuentes[size] = tamanos_fuentes.get(size, 0) + len(span["text"])
                        
    fuente_base = max(tamanos_fuentes, key=tamanos_fuentes.get) if tamanos_fuentes else 10.0

    for num_pag, pagina in enumerate(doc, start=1):
        alto_pagina = pagina.rect.height
        bloques = pagina.get_text("dict")["blocks"]
        
        for b in bloques:
            if b["type"] == 0:
                y0 = b["bbox"][1]
                if y0 < alto_pagina * 0.08 or y0 > alto_pagina * 0.92:
                    continue

                texto_bloque = ""
                max_size_in_block = 0
                for linea in b["lines"]:
                    for span in linea["spans"]:
                        texto_bloque += span["text"] + " "
                        if span["size"] > max_size_in_block:
                            max_size_in_block = round(span["size"], 1)
                
                texto_bloque = texto_bloque.strip()
                if not texto_bloque: continue
                
                if max_size_in_block > fuente_base + 3: tipo = 'titulo_capitulo'
                elif max_size_in_block > fuente_base + 1: tipo = 'titulo_seccion'
                else: tipo = 'texto'
                    
                bloques_estructurados.append((tipo, texto_bloque, num_pag))
                
    return bloques_estructurados

# ============================================================
# MOTOR DE TRADUCCIÓN RESILIENTE
# ============================================================
def dividir_texto(texto, max_len=MAX_CHARS_PER_FRAGMENT):
    if len(texto) <= max_len: return [texto]
    fragmentos = []
    while len(texto) > max_len:
        idx = texto.rfind('. ', 0, max_len)
        if idx == -1 or idx < max_len * 0.5: idx = texto.rfind(' ', 0, max_len)
        if idx == -1: idx = max_len
        fragmentos.append(texto[:idx+1])
        texto = texto[idx+1:].lstrip()
    if texto: fragmentos.append(texto)
    return fragmentos

def traducir_bloque(idx_bloque, bloque, idioma_destino):
    tipo, contenido, pagina = bloque
    
    if not necesita_traduccion(contenido):
        return (idx_bloque, tipo, contenido, pagina, {"reqs": 0, "reintentos": 0, "fallbacks": 0})
        
    traductor = GoogleTranslator(source='auto', target=idioma_destino)
    fragmentos = dividir_texto(contenido)
    traducciones = []
    stats = {"reqs": 0, "reintentos": 0, "fallbacks": 0}
    
    for frag in fragmentos:
        intentos = 0
        exito = False
        while intentos < 3 and not exito:
            stats["reqs"] += 1
            try:
                trad = traductor.translate(frag)
                if trad and ("Error 500" in trad or "Server Error" in trad):
                    raise Exception("Falso positivo API")
                traducciones.append(trad)
                exito = True
                time.sleep(0.3) 
            except Exception:
                intentos += 1
                stats["reintentos"] += 1
                if intentos < 3:
                    time.sleep(3 * intentos) 
                    
        if not exito:
            stats["fallbacks"] += 1
            traducciones.append(frag) 
            
    return (idx_bloque, tipo, ' '.join(traducciones), pagina, stats)

def procesar_pipeline(bloques, file_hash, ui_metrics, idioma_destino='es'):
    chk = cargar_checkpoint(file_hash)
    resultados_dict = chk["resultados"]
    metricas = chk["metricas"]
    
    total_bloques = len(bloques)
    pendientes = [i for i in range(total_bloques) if str(i) not in resultados_dict]
    
    if not pendientes:
        return [resultados_dict[str(i)] for i in range(total_bloques)], metricas
        
    lotes = [pendientes[i:i + TAMANO_LOTE] for i in range(0, len(pendientes), TAMANO_LOTE)]
    pausa_actual = PAUSA_BASE
    
    for idx_lote, lote_indices in enumerate(lotes):
        sub_resultados = {}
        errores_lote = 0
        t0 = time.time()
        
        with ThreadPoolExecutor(max_workers=WORKERS_PARALELOS) as executor:
            futuros = [executor.submit(traducir_bloque, idx, bloques[idx], idioma_destino) for idx in lote_indices]
            
            for futuro in as_completed(futuros):
                idx, tipo, trad, pag, stats = futuro.result()
                sub_resultados[str(idx)] = (tipo, trad, pag)
                
                metricas["requests_enviados"] += stats["reqs"]
                metricas["reintentos"] += stats["reintentos"]
                metricas["fallbacks"] += stats["fallbacks"]
                errores_lote += stats["reintentos"]
        
        metricas["tiempo_acumulado"] += (time.time() - t0)
        resultados_dict.update(sub_resultados)
        chk["procesados"] = len(resultados_dict)
        
        guardar_checkpoint(chk, file_hash)
        
        if errores_lote == 0:
            pausa_actual = max(PAUSA_BASE, pausa_actual - 0.5)
        elif errores_lote < 3:
            pausa_actual = min(MAX_PAUSA, pausa_actual + 2.0)
        else:
            pausa_actual = min(MAX_PAUSA, pausa_actual + 5.0)
            
        tasa_fallbacks = (metricas["fallbacks"] / max(1, metricas["requests_enviados"])) * 100
        
        with ui_metrics.container():
            st.progress(chk["procesados"] / total_bloques)
            c1, c2, c3, c4 = st.columns(4)
            c1.metric("Procesados", f"{chk['procesados']}/{total_bloques}")
            c2.metric("Peticiones API", metricas["requests_enviados"])
            c3.metric("Tasa Fallback", f"{tasa_fallbacks:.2f}%", f"{metricas['fallbacks']} fallos", delta_color="inverse")
            c4.metric("Pausa Dinámica", f"{pausa_actual:.1f}s", "Anti-Ban activo")

        if idx_lote < len(lotes) - 1:
            time.sleep(pausa_actual)

    lista_final_ordenada = [resultados_dict[str(i)] for i in range(total_bloques)]
    return lista_final_ordenada, metricas

# ============================================================
# EXPORTACIÓN A WORD (.DOCX)
# ============================================================
def generar_word(resultados):
    documento = Document()
    for tipo, contenido, _ in resultados:
        if not contenido or not contenido.strip():
            continue
            
        if tipo == "titulo_capitulo":
            parrafo = documento.add_paragraph(contenido, style="Heading 1")
        elif tipo == "titulo_seccion":
            parrafo = documento.add_paragraph(contenido, style="Heading 2")
        else:
            parrafo = documento.add_paragraph(contenido, style="Normal")
            parrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
    buffer = io.BytesIO()
    documento.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# ============================================================
# EXPORTACIÓN A PDF (FORMATO LIBRO)
# ============================================================
def generar_pdf_libro(resultados, nombre_archivo="Documento_Traducido.pdf"):
    if not os.path.exists(FONT_PATH):
        st.error(f"Falta fuente en: {FONT_PATH}")
        st.stop()

    if "DejaVuSans" not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont("DejaVuSans", FONT_PATH))

    buffer = io.BytesIO()
    doc = BaseDocTemplate(
        buffer, pagesize=A4, leftMargin=25*mm, rightMargin=20*mm, 
        topMargin=25*mm, bottomMargin=25*mm, title="Traducción Académica", 
        author="Traductor Académico Industrial"
    )
    ancho, alto = A4

    frame_body = Frame(25*mm, 25*mm, ancho-45*mm, alto-50*mm, id="body")
    frame_cover = Frame(0, 0, ancho, alto, leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0, id="cover")
    frame_toc = Frame(25*mm, 25*mm, ancho-45*mm, alto-50*mm, id="toc")

    estilos = getSampleStyleSheet()
    e_norm = ParagraphStyle("LibroNormal", parent=estilos["Normal"], fontName="DejaVuSans", fontSize=11, leading=15, alignment=TA_JUSTIFY, spaceAfter=10, firstLineIndent=5*mm)
    e_cap = ParagraphStyle("LibroCapitulo", parent=estilos["Heading1"], fontName="DejaVuSans", fontSize=16, leading=20, alignment=TA_CENTER, spaceBefore=20, spaceAfter=15, keepWithNext=True)
    e_sec = ParagraphStyle("LibroSeccion", parent=estilos["Heading2"], fontName="DejaVuSans", fontSize=13, leading=16, alignment=TA_LEFT, spaceBefore=12, spaceAfter=10, keepWithNext=True)
    e_toc_title = ParagraphStyle("TOCTitulo", parent=estilos["Title"], fontName="DejaVuSans", fontSize=20, leading=24, alignment=TA_CENTER, spaceAfter=20)
    e_toc_cap = ParagraphStyle("TOCCapitulo", parent=estilos["Normal"], fontName="DejaVuSans", fontSize=11, leading=15, leftIndent=0, firstLineIndent=0, spaceAfter=5)
    e_toc_sec = ParagraphStyle("TOCSeccion", parent=estilos["Normal"], fontName="DejaVuSans", fontSize=10, leading=14, leftIndent=10*mm, firstLineIndent=0, spaceAfter=3)

    def dibujar_portada(canvas, doc):
        canvas.saveState()
        titulo = Paragraph("Traducción Académica", ParagraphStyle("PortadaTitulo", fontName="DejaVuSans", fontSize=26, leading=32, alignment=TA_CENTER))
        subtitulo = Paragraph(escape(nombre_archivo), ParagraphStyle("PortadaArchivo", fontName="DejaVuSans", fontSize=13, leading=18, alignment=TA_CENTER))
        _, h_titulo = titulo.wrap(ancho - 40*mm, alto)
        _, h_subtitulo = subtitulo.wrap(ancho - 40*mm, alto)
        espacio_entre = 10 * mm
        alto_total = h_titulo + espacio_entre + h_subtitulo
        y_inicio = (alto - alto_total) / 2
        titulo.drawOn(canvas, 20*mm, y_inicio + h_subtitulo + espacio_entre)
        subtitulo.drawOn(canvas, 20*mm, y_inicio)
        canvas.restoreState()

    def dibujar_cuerpo(canvas, doc):
        canvas.saveState()
        pagina = canvas.getPageNumber()
        canvas.setLineWidth(0.5)
        canvas.line(25*mm, alto - 18*mm, ancho - 20*mm, alto - 18*mm)
        canvas.setFont("DejaVuSans", 8)
        canvas.drawString(25*mm, alto - 14*mm, "TRADUCCIÓN ACADÉMICA")
        canvas.drawRightString(ancho - 20*mm, alto - 14*mm, "Edición académica")
        canvas.line(25*mm, 19*mm, ancho - 20*mm, 19*mm)
        canvas.setFont("DejaVuSans", 9)
        canvas.drawCentredString(ancho / 2, 12*mm, str(pagina))
        canvas.restoreState()

    def dibujar_toc(canvas, doc):
        canvas.saveState()
        canvas.setFont("DejaVuSans", 8)
        canvas.drawCentredString(ancho / 2, 12*mm, str(canvas.getPageNumber()))
        canvas.restoreState()

    doc.addPageTemplates([
        PageTemplate(id="PORTADA", frames=[frame_cover], onPage=dibujar_portada),
        PageTemplate(id="TOC", frames=[frame_toc], onPage=dibujar_toc),
        PageTemplate(id="CUERPO", frames=[frame_body], onPage=dibujar_cuerpo)
    ])

    toc = TableOfContents()
    toc.levelStyles = [
        ParagraphStyle("TOCLevel1", fontName="DejaVuSans", fontSize=11, leading=15, leftIndent=0, firstLineIndent=0, spaceBefore=5, spaceAfter=3),
        ParagraphStyle("TOCLevel2", fontName="DejaVuSans", fontSize=10, leading=14, leftIndent=10*mm, firstLineIndent=0, spaceBefore=2, spaceAfter=2)
    ]

    contador_titulos = [0]
    
    def after_flowable(flowable):
        if not isinstance(flowable, Paragraph): return
        estilo = flowable.style.name
        if estilo == "LibroCapitulo": nivel = 0
        elif estilo == "LibroSeccion": nivel = 1
        else: return
        
        texto = flowable.getPlainText()
        contador_titulos[0] += 1
        key = f"heading_{contador_titulos[0]}"
        
        try:
            canvas = doc.canv
            canvas.bookmarkPage(key)
            canvas.addOutlineEntry(texto, key, level=nivel, closed=False)
        except Exception:
            pass
        doc.notify("TOCEntry", (nivel, texto, doc.page, key))
        
    doc.afterFlowable = after_flowable

    historia = []
    historia.append(NextPageTemplate("TOC"))
    historia.append(PageBreak())
    historia.append(Paragraph("Índice", e_toc_title))
    historia.append(toc)
    historia.append(NextPageTemplate("CUERPO"))
    historia.append(PageBreak())

    for tipo, contenido, _ in resultados:
        if not contenido or not contenido.strip(): continue
        contenido_escape = escape(contenido.strip())
        if tipo == "titulo_capitulo": historia.append(Paragraph(contenido_escape, e_cap))
        elif tipo == "titulo_seccion": historia.append(Paragraph(contenido_escape, e_sec))
        else: historia.append(Paragraph(contenido_escape, e_norm))

    doc.multiBuild(historia)
    buffer.seek(0)
    return buffer.getvalue()

# ============================================================
# LOG DE OPERACIÓN
# ============================================================
def generar_log(metricas, total_bloques):
    lineas = [
        "="*60, "LOG TELEMÉTRICO - TRADUCCIÓN INDUSTRIAL",
        f"Fecha finalización: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        f"Tiempo total activo (s): {metricas['tiempo_acumulado']:.2f}",
        "-"*60,
        f"Bloques Totales: {total_bloques}",
        f"Peticiones enviadas a la API: {metricas['requests_enviados']}",
        f"Reintentos ejecutados: {metricas['reintentos']}",
        f"Fallbacks (Texto original forzado por error de red): {metricas['fallbacks']}",
        "="*60
    ]
    return "\n".join(lineas)

# ============================================================
# INTERFAZ STREAMLIT
# ============================================================
st.set_page_config(page_title="Traductor Académico (Industrial)", layout="centered")
st.title("📚 Traductor de PDFs Académicos (Industrial)")
st.markdown("Equipado con **Checkpoints atómicos en disco**, telemetría en tiempo real, exportación en formato Libro y formato Word nativo.")

if "traduccion_lista" not in st.session_state: st.session_state.traduccion_lista = False
if "pdf_final" not in st.session_state: st.session_state.pdf_final = None
if "word_final" not in st.session_state: st.session_state.word_final = None
if "log_final" not in st.session_state: st.session_state.log_final = None

archivo = st.file_uploader("Sube el PDF masivo", type="pdf")

if archivo is not None:
    bytes_pdf = archivo.read()
    hash_pdf = generar_hash_archivo(bytes_pdf)
    
    if "hash_actual" not in st.session_state or st.session_state.hash_actual != hash_pdf:
        with st.spinner("Mapeando topología del PDF..."):
            st.session_state.bloques = detectar_estructura_pymupdf(bytes_pdf)
            st.session_state.hash_actual = hash_pdf
            st.session_state.traduccion_lista = False
            
            chk = cargar_checkpoint(hash_pdf)
            st.session_state.progreso_previo = chk["procesados"]
            
    total = len(st.session_state.bloques)
    procesados = st.session_state.progreso_previo
    
    if 0 < procesados < total:
        st.warning(f"🔄 Checkpoint detectado. Se reanudará desde el bloque {procesados}/{total}.")
    elif procesados == total:
        st.success(f"✅ Documento previamente procesado al 100%.")
    else:
        st.info(f"📄 Nuevo documento analizado: {total} bloques listos para procesar.")

    ui_metrics = st.empty()

    if st.button("🚀 Procesar", type="primary", use_container_width=True):
        resultados, metricas = procesar_pipeline(st.session_state.bloques, hash_pdf, ui_metrics)
        
        with st.spinner("Generando libro académico, Word editable y log..."):
            st.session_state.pdf_final = generar_pdf_libro(resultados, nombre_archivo=archivo.name)
            st.session_state.word_final = generar_word(resultados)
            st.session_state.log_final = generar_log(metricas, total)
            st.session_state.traduccion_lista = True
            
        st.rerun()

if st.session_state.traduccion_lista:
    st.markdown("---")
    st.success("🎉 Documento ensamblado exitosamente.")
    
    c1, c2, c3 = st.columns(3)
    with c1:
        st.download_button(
            "📥 Descargar PDF (Formato Libro)", data=st.session_state.pdf_final,
            file_name="Libro_Traducido.pdf", mime="application/pdf", use_container_width=True
        )
    with c2:
        st.download_button(
            "📝 Descargar Word (.docx editable)", data=st.session_state.word_final,
            file_name="Libro_Traducido.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
    with c3:
        st.download_button(
            "📋 Descargar Log de Calidad", data=st.session_state.log_final,
            file_name="Log_Operacion.txt", mime="text/plain", use_container_width=True
        )
